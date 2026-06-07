import os
from docx import Document

class Exporter:
    def __init__(self):
        pass

    def export(self, title: str, sections: list, output_dir: str, format: str) -> str:
        os.makedirs(output_dir, exist_ok=True)
        if format.lower() == 'markdown' or format.lower() == 'md':
            return self.export_markdown(title, sections, output_dir)
        elif format.lower() == 'latex' or format.lower() == 'tex' or format.lower() == 'zip':
            return self.export_latex(title, sections, output_dir)
        elif format.lower() == 'pdf':
            # Basic fallback for PDF since no library is present, we return markdown text
            return self.export_markdown(title, sections, output_dir, as_pdf=True)
        else:
            return self.export_docx(title, sections, output_dir)

    def export_markdown(self, title: str, sections: list, output_dir: str, as_pdf: bool = False) -> str:
        filename = f"{title.replace(' ', '_')}.{'pdf' if as_pdf else 'md'}"
        filepath = os.path.join(output_dir, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"# {title}\n\n")
            for sec in sections:
                if sec.get('heading'):
                    f.write(f"## {sec.get('heading')}\n\n")
                if sec.get('content'):
                    f.write(f"{sec.get('content')}\n\n")
        return filepath
        
    def export_latex(self, title: str, sections: list, output_dir: str) -> str:
        filename = f"{title.replace(' ', '_')}.zip" # using zip/tex extension to match frontend
        filepath = os.path.join(output_dir, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("\\documentclass{article}\n\\begin{document}\n")
            f.write(f"\\title{{{title}}}\n\\maketitle\n\n")
            for sec in sections:
                if sec.get('heading'):
                    f.write(f"\\section{{{sec.get('heading')}}}\n")
                if sec.get('content'):
                    f.write(f"{sec.get('content')}\n\n")
            f.write("\\end{document}\n")
        return filepath

    def export_docx(self, title: str, sections: list, output_dir: str) -> str:
        """
        Exports the verified draft to a DOCX file.
        sections should be a list of dicts: [{"heading": "Introduction", "content": "..."}]
        """
        os.makedirs(output_dir, exist_ok=True)
        filename = f"{title.replace(' ', '_')}.docx"
        filepath = os.path.join(output_dir, filename)
        
        doc = Document()
        doc.add_heading(title, 0)
        
        for section in sections:
            doc.add_heading(section.get('heading', ''), level=1)
            doc.add_paragraph(section.get('content', ''))
            
        doc.save(filepath)
        return filepath

exporter_engine = Exporter()
